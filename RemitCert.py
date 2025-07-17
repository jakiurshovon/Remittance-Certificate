import os
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from datetime import datetime
import math

def get_clean_number(value):
    try:
        if pd.isna(value):
            return 0.0
        cleaned = str(value).replace(',', '').replace('\xa0', '').strip()
        return float(cleaned) if cleaned else 0.0
    except:
        return 0.0

def is_incentive_match(actual, expected, tolerance=0.1):
    return (
        abs(actual - expected) <= tolerance or
        abs(actual - round(expected)) <= tolerance or
        abs(actual - math.floor(expected)) <= tolerance or
        abs(actual - math.ceil(expected)) <= tolerance
    )

def convert_number_to_words(n):
    units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
    teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
             "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

    def two_digits(num):
        if num < 10:
            return units[num]
        elif 10 <= num < 20:
            return teens[num - 10]
        else:
            return tens[num // 10] + (" " + units[num % 10] if num % 10 else "")

    def three_digits(num):
        h = num // 100
        rest = num % 100
        result = ""
        if h:
            result += units[h] + " Hundred"
        if rest:
            result += (" and " if result else "") + two_digits(rest)
        return result

    integer_part = int(n)
    decimal_part = int(round((n - integer_part) * 100))

    crore = integer_part // 10000000
    lac = (integer_part % 10000000) // 100000
    thousand = (integer_part % 100000) // 1000
    below_thousand = integer_part % 1000

    parts = []
    if crore: parts.append(two_digits(crore) + " Crore")
    if lac: parts.append(two_digits(lac) + " Lac")
    if thousand: parts.append(two_digits(thousand) + " Thousand")
    if below_thousand: parts.append(three_digits(below_thousand))

    taka_words = "Taka " + " ".join(parts) if parts else "Taka Zero"

    if decimal_part > 0:
        paisa_words = " and Paisa " + two_digits(decimal_part)
    else:
        paisa_words = ""

    return taka_words + paisa_words + " Only"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def insert_table_after_paragraph(paragraph, document, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    paragraph._element.addnext(table._tbl)
    return table

def insert_paragraph_after(paragraph, text):
    new_par = paragraph.insert_paragraph_before(text)
    paragraph._element.addnext(new_par._element)
    return new_par

def generate_certificate(filepath):
    try:
        xl = pd.ExcelFile(filepath)
        df_data = xl.parse(xl.sheet_names[0])
        keyword_file = os.path.join("data", "Remittance Certificate Key Words.xlsx")
        kw_xl = pd.ExcelFile(keyword_file)
        principal_keywords = kw_xl.parse('Principal').iloc[:, 0].dropna().astype(str).tolist()
        incentive_keywords = kw_xl.parse('Incentive').iloc[:, 0].dropna().astype(str).tolist()
        exclusion_keywords = kw_xl.parse('Exclusion').iloc[:, 0].dropna().astype(str).str.upper().tolist()

        template = os.path.join("template", "Template.docx")
        doc = Document(template)
        customer_name = df_data.iloc[6, 0]
        account_no = df_data.iloc[6, 1]
        date_period = df_data.iloc[7, 1]
        branch = df_data.iloc[1, 0]

        for para in doc.paragraphs:
            if "Mr." in para.text:
                para.text = para.text.replace("Mr.", f"Mr./Ms. {customer_name}")
            if "A/C No-" in para.text:
                para.text = para.text.replace("A/C No-", f" {account_no}")
            if "period" in para.text:
                para.text = para.text.replace("period", f" {date_period}")
            if "PLC," in para.text:
                para.text = para.text.replace("PLC,", f"PLC, {branch}")
            if "Ref:" in para.text:
                insert_paragraph_after(para, f"{datetime.today().strftime('%B %d, %Y')}")

        insert_after = next((p for p in doc.paragraphs if "Table-A:  Remittance Details:" in p.text), None)

        # Step 1: Collect all eligible rows
        credit_rows = []
        for i in range(1, len(df_data)):
            credit = get_clean_number(df_data.iloc[i, 5])
            if credit <= 0:
                continue

            date = df_data.iloc[i, 0]
            brn_raw = df_data.iloc[i, 1]
            desc = str(df_data.iloc[i, 2])

            # Check if date or BRN is valid
            try:
                brn = str(int(float(brn_raw))).zfill(3)
            except:
                brn = ""

            try:
                pd.to_datetime(date)
                valid_date = True
            except:
                valid_date = False

            if not brn.isdigit() and not valid_date:
                continue

            credit_rows.append({
                "index": i,
                "date": date,
                "brn": brn,
                "desc": desc,
                "desc_upper": desc.upper(),
                "credit": credit
            })

        # Step 2: First pass - detect all incentive amounts and mark them
        used_indices = set()
        incentive_indices = set()
        tolerance = 0.1
        potential_incentives = []

        # Pass 1: Identify all likely incentive rows
        for i, row in enumerate(credit_rows):
            for j, other in enumerate(credit_rows):
                if i == j or other["index"] in used_indices or row["index"] in used_indices:
                    continue

                expected_incentives = [row["credit"] * 0.025, row["credit"] * 0.02]

                # ✅ Check if BR code rules match
                if row["brn"] == "747" and other["brn"] != "747":
                    continue  # skip if principal is 747 but incentive is not
                # (No check for BRN 100 — it's allowed as-is)

                if any(is_incentive_match(other["credit"], ei, tolerance) for ei in expected_incentives):
                    potential_incentives.append((row["index"], other["index"]))
                    used_indices.add(other["index"])
                    incentive_indices.add(other["index"])
                    break

        # Step 2: Principal matching
        principal_rows = []
        for row in credit_rows:
            i = row["index"]
            brn = row["brn"]
            desc = row["desc"]
            desc_upper = row["desc_upper"]
            credit = row["credit"]

            if i in incentive_indices:
                continue  # Skip incentive rows

            # ⛔ Skip descriptions with exclusion keywords
            if any(ex_kw in desc_upper for ex_kw in exclusion_keywords):
                continue

            matched_incentive = 0.0
            for principal_idx, incentive_idx in potential_incentives:
                if principal_idx == i:
                    matched_incentive = get_clean_number(df_data.iloc[incentive_idx, 5])
                    used_indices.add(i)
                    used_indices.add(incentive_idx)
                    break

            # ... rest of your existing matching logic for BRN 747 and 100 ...

            if brn == "747" and any(k.upper() in desc_upper for k in principal_keywords):
                principal_rows.append({
                    "index": i,
                    "desc": "EFT Credited by Other Bank",
                    "amount": credit,
                    "matched_incentive": matched_incentive
                })
            elif brn == "100" and ("NPSB" in desc_upper or any(k.upper() in desc_upper for k in principal_keywords)):
                is_npsb = "NPSB" in desc_upper
                if is_npsb:
                    if matched_incentive > 0:
                        principal_rows.append({
                            "index": i,
                            "desc": desc,
                            "amount": credit,
                            "matched_incentive": matched_incentive
                        })
                else:
                    principal_rows.append({
                        "index": i,
                        "desc": desc,
                        "amount": credit,
                        "matched_incentive": matched_incentive
                    })

        # Step 3: Build table
        table = insert_table_after_paragraph(insert_after, doc, 1, 6) if insert_after else doc.add_table(rows=1, cols=6)
        try:
            table.style = 'Table Grid'
        except:
            pass
        set_table_borders(table)

        hdr = table.rows[0].cells
        headers = ["Sl.", "Credit Date", "Reference/Description", "Principal Amount (BDT)", "Incentive Amount (BDT)", "Total Amount (BDT)"]
        for i in range(6):
            hdr[i].text = headers[i]
            hdr[i].paragraphs[0].runs[0].font.bold = True

        column_widths = [Inches(0.3), Inches(1.4), Inches(2.4), Inches(1.2), Inches(1.2), Inches(1.4)]
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = width

        total_p = total_i = 0
        for idx, row in enumerate(principal_rows, 1):
            i = row["index"]
            date = df_data.iloc[i, 0]
            desc = row["desc"]
            principal_amt = row["amount"]
            incentive_amt = row["matched_incentive"]
            total_amt = principal_amt + incentive_amt
            total_p += principal_amt
            total_i += incentive_amt

            r = table.add_row().cells
            r[0].text = str(idx)
            r[1].text = date.strftime('%B %d, %Y') if isinstance(date, datetime) else str(date)
            r[2].text = desc
            r[3].text = f"{principal_amt:,.2f}"
            r[4].text = f"{incentive_amt:,.2f}"
            r[5].text = f"{total_amt:,.2f}"

        total_row = table.add_row().cells
        totals = ["", "", "Total", f"{total_p:,.2f}", f"{total_i:,.2f}", f"{total_p + total_i:,.2f}"]
        for i, val in enumerate(totals):
            run = total_row[i].paragraphs[0].add_run(val)
            run.bold = True

        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        word_row = table.add_row().cells
        merged = word_row[0].merge(word_row[5])
        para = merged.paragraphs[0]
        run = para.add_run("In Words: " + convert_number_to_words(total_p + total_i))
        run.bold = True
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        os.makedirs("output", exist_ok=True)
        out_path = os.path.join("output", f"Remittance_Certificate_{customer_name.replace(' ', '_')}.docx")
        doc.save(out_path)
        messagebox.showinfo("Success", f"Certificate saved to:\n{out_path}")

    except Exception as e:
        messagebox.showerror("Error", str(e))

def browse_file():
    filename = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx *.xls")])
    if filename:
        file_path.set(filename)

def start_generate():
    if not file_path.get():
        messagebox.showwarning("No file", "Please select an Excel statement file.")
        return
    generate_certificate(file_path.get())

app = tk.Tk()
app.title("Remittance Certificate Generator")
app.geometry("520x220")

file_path = tk.StringVar()
tk.Label(app, text="Browse Excel Statement:").pack(pady=10)
tk.Entry(app, textvariable=file_path, width=60).pack()
tk.Button(app, text="Browse", command=browse_file).pack(pady=5)
tk.Button(app, text="Generate Certificate", command=start_generate, bg="green", fg="white").pack(pady=20)

app.mainloop()

